# -*- coding: utf-8 -*-
"""
Extracts the 2024-2025 laureates list from REMISE DES PRIX 2025-0.1.docx
into a clean, structured JSON ready for import into Supabase.

Run with: PYTHONUTF8=1 python scripts/extract_laureates_2025.py
"""

import difflib
import docx
import json
import re
import unicodedata

SOURCE = "REMISE DES PRIX 2025-0.1.docx"
OUTPUT = "data/laureates-2024-2025.json"

ANGLO_MARKERS = re.compile(r"CLASS|FORM|FROM|LOWER|UPPER|YEAR", re.IGNORECASE)

# Raw class label -> canonical niveau code (matching critères.json's taxonomy).
NIVEAU_MAP = {
    "": None,
    "PN": "JARDIN",
    "PS": "JARDIN",
    "MS": "JARDIN",
    "GS": "JARDIN",
    "N1": "JARDIN",
    "N2": "JARDIN",
    "SIL": "SIL",
    "CP": "CP",
    "CE1": "CE1",
    "CE2": "CE2",
    "CM1": "CM1",
    "CM2": "CM2",
    "6E": "6e",
    "5E": "5e",
    "4E": "4e",
    "3E": "3e",
    "2ND": "2nd",
    "2NDC": "2nd",
    "1ER": "1ere",
    "1ERE": "1ere",
    "1ERE TI": "1ere",
    "TLE": "Tle",
    "TLEC": "Tle",
    "TLEAESP": "Tle",
    "1ER ANNEE": "UNIVERSITE",
    "1ERE ANNEE": "UNIVERSITE",
    "NIV 1": "UNIVERSITE",
    "NIV 2": "UNIVERSITE",
    "NIV2": "UNIVERSITE",
    "NIV 3": "UNIVERSITE",
    "NIV3": "UNIVERSITE",
    "NIV4": "UNIVERSITE",
    "CLASS 1": "CLASS 1",
    "CLASS 2": "CLASS 2",
    "CLASS 3": "CLASS 3",
    "CLASS 4": "CLASS 4",
    "CLASS 5": "CLASS 5",
    "CLASS 6": "CLASS 6",
    "FROM 1": "FORM 1",
    "FROM 2": "FORM 2",
    "FROM 3": "FORM 3",
    "FROM 5": "FORM 5",
    "LOWER 6": "LOWER 6",
    "LOWER6": "LOWER 6",
    "LOWER 5": None,  # not a recognized level — flagged for manual review
    "UPPER 6": "UPPER 6",
    "YEAR 1": "UNIVERSITY",
}


def guess_section(depart: str, arrivee: str) -> str:
    if ANGLO_MARKERS.search(depart) or ANGLO_MARKERS.search(arrivee):
        return "anglophone"
    return "francophone"


def strip_accents(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c))


def normalize_niveau(raw: str):
    key = strip_accents(raw.strip()).upper()
    if key in NIVEAU_MAP:
        return NIVEAU_MAP[key]
    return None  # unmapped, flagged for manual review


def normalize_name(name: str) -> str:
    name = strip_accents(name)
    name = re.sub(r"\s+", " ", name).strip().upper()
    return name


def parse_moyenne(raw: str):
    raw = raw.strip().replace(",", ".")
    if not raw:
        return None
    try:
        return float(raw)
    except ValueError:
        return None


def parse_rang(observation: str):
    m = re.match(r"\s*(\d+)\s*e", observation, re.IGNORECASE)
    return int(m.group(1)) if m else None


def extract_table(tbl):
    rows = []
    for row in tbl.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def main():
    d = docx.Document(SOURCE)
    special_rows = extract_table(d.tables[0])
    excellence_rows = extract_table(d.tables[1])
    encouragement_rows = extract_table(d.tables[2])
    excellence_plus_rows = extract_table(d.tables[3])

    laureates = {}  # normalized_name -> record
    unmapped_niveaux = set()
    ambiguous_sections = []
    duplicate_names = []

    def add(
        name,
        section,
        depart_raw,
        arrivee_raw,
        moyenne,
        rang,
        prize,
        observation="",
    ):
        key = normalize_name(name)
        depart = normalize_niveau(depart_raw)
        arrivee = normalize_niveau(arrivee_raw) if arrivee_raw else None
        if depart_raw and depart is None:
            unmapped_niveaux.add(depart_raw)
        if arrivee_raw and arrivee is None:
            unmapped_niveaux.add(arrivee_raw)

        # Some graduate-level distinctions (e.g. "Master en traduction") have
        # no class fields at all in the source — treat as a university-level
        # terminal achievement rather than dropping the niveau_depart entirely
        # (results.niveau_depart is NOT NULL).
        if depart is None and not depart_raw and not arrivee_raw:
            depart = "UNIVERSITE"
            ambiguous_sections.append(
                (name, f"(vide, déduit de: {observation})", "")
            )

        # University-level rows have no anglo/francophone marker in the raw
        # label, so the section guess (defaulting to francophone) is unverified.
        if depart == "UNIVERSITE" and not ANGLO_MARKERS.search(
            depart_raw + arrivee_raw
        ):
            ambiguous_sections.append((name, depart_raw, arrivee_raw))

        if key in laureates:
            duplicate_names.append(name)
            laureates[key]["prizes"].append(prize)
            return

        laureates[key] = {
            "name": name,
            "section": section,
            "niveau_depart": depart,
            "niveau_admission": arrivee,
            "moyenne": moyenne,
            "rang": rang,
            "prizes": [prize],
        }

    for row in special_rows:
        _, name, depart, arrivee = row
        section = guess_section(depart, arrivee)
        add(name, section, depart, arrivee, None, None, "SPECIAL")

    for row in excellence_rows:
        _, name, depart, arrivee, moyenne, obs = row
        section = guess_section(depart, arrivee)
        add(
            name,
            section,
            depart,
            arrivee,
            parse_moyenne(moyenne),
            parse_rang(obs),
            "EXC",
            obs,
        )

    for row in encouragement_rows:
        _, name, depart, arrivee, moyenne, obs = row
        section = guess_section(depart, arrivee)
        add(
            name,
            section,
            depart,
            arrivee,
            parse_moyenne(moyenne),
            parse_rang(obs),
            "ENC",
            obs,
        )

    # EXCELLENCE_PLUS: cross-reference by normalized name against EXCELLENCE list.
    # Falls back to a close-match suggestion (not auto-applied) when the exact
    # normalized name isn't found, since the source document has some typos
    # between the EXCELLENCE table and the EXCELLENCE_PLUS table.
    unmatched_plus = []
    fuzzy_matched_plus = []
    exc_keys = [k for k, v in laureates.items() if "EXC" in v["prizes"]]
    for row in excellence_plus_rows:
        _, name = row
        key = normalize_name(name)
        if key in laureates and "EXC" in laureates[key]["prizes"]:
            laureates[key]["prizes"].append("EXC_PLUS")
            continue

        suggestions = difflib.get_close_matches(
            key, exc_keys, n=2, cutoff=0.75
        )
        if len(suggestions) == 1:
            # Single confident match — the source doc has a spelling typo
            # between the EXCELLENCE and EXCELLENCE_PLUS tables for this row.
            matched_key = suggestions[0]
            laureates[matched_key]["prizes"].append("EXC_PLUS")
            fuzzy_matched_plus.append(
                {
                    "listed_as": name,
                    "matched_to": laureates[matched_key]["name"],
                }
            )
        else:
            unmatched_plus.append(
                {
                    "name": name,
                    "suggested_matches": [
                        laureates[s]["name"] for s in suggestions
                    ],
                }
            )

    result = {
        "school_year": "2024-2025",
        "start_year": 2024,
        "laureates": list(laureates.values()),
        "counts": {
            "special": len(special_rows),
            "excellence": len(excellence_rows),
            "encouragement": len(encouragement_rows),
            "excellence_plus_listed": len(excellence_plus_rows),
            "excellence_plus_matched": sum(
                1 for l in laureates.values() if "EXC_PLUS" in l["prizes"]
            ),
            "unique_students": len(laureates),
        },
        "warnings": {
            "unmapped_niveaux": sorted(unmapped_niveaux),
            "ambiguous_sections": ambiguous_sections,
            "duplicate_names_within_docx": duplicate_names,
            "excellence_plus_fuzzy_matched": fuzzy_matched_plus,
            "excellence_plus_unmatched": unmatched_plus,
        },
    }

    with open(OUTPUT, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(json.dumps(result["counts"], indent=2, ensure_ascii=False))
    print("\nWarnings:")
    print(json.dumps(result["warnings"], indent=2, ensure_ascii=False))
    print(f"\nSaved to {OUTPUT}")


if __name__ == "__main__":
    main()
